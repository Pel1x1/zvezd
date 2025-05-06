from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches

def create_formatted_document():
    # Create a new Document
    doc = Document()
    
    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1)
    
    # Add first paragraph
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p1.paragraph_format.left_indent = Cm(0.5)
    p1.paragraph_format.first_line_indent = Cm(1.5)
    p1.paragraph_format.line_spacing = 1.5
    run = p1.add_run("Microsoft Word - это одна из главных программ пакета MS Office, предназначенная для создания, просмотра и печати текстовых документов. Word - одна из самых совершенных программ в сфере текстовых процессоров, которая предполагает выполнение всех основных операций над текстовой информацией. С помощью Word можно быстро и с высоким качеством подготовить любой документ - от простой записи до оригиналь-макета цветного издания.")
    run.font.name = 'Courier New'
    run.font.size = Pt(13)
    run.font.italic = True
    
    # Add table
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    # Fill table data
    data = [
        ['8', '6', '7'],
        ['6', '4', '5'],
        ['3', '2', '1']
    ]
    for i, row in enumerate(data):
        for j, value in enumerate(row):
            table.cell(i, j).text = value
    
    # Add second paragraph
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p2.paragraph_format.right_indent = Cm(0.5)
    p2.paragraph_format.first_line_indent = Cm(2.0)
    run = p2.add_run("Во-первых, Word даёт возможность выполнять все базовые редакционные операции над текстом, предусмотренные современной компьютерной технологией.")
    run.font.name = 'Arial'
    run.font.size = Pt(12)
    
    # Add third paragraph
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p3.paragraph_format.first_line_indent = Cm(1.0)
    p3.paragraph_format.space_after = Pt(10)
    run = p3.add_run("Во-вторых, в процессоре Word реализованы возможности новейших технологий связывания объектов OLE, которая позволяет включать в документ фрагменты, таблицы, изображения, подготовленные в других приложениях Windows. Встроенные объекты можно редактировать средствами этих приложений.")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    
    # Add signature section
    p_sig = doc.add_paragraph()
    p_sig.add_run("Старший преподаватель\nкафедры информационного обеспечения ОВА\nПодписник Политов\t\t\tП.Б. Сергеев\n22.12.2022")
    
    # Save the document
    doc.save('formatted_document.docx')

if __name__ == "__main__":
    create_formatted_document() 